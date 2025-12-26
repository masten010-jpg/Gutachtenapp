# programm_2_word_output.py  # Kommentar: Programm 2 – KI JSON -> Word, inkl. Summenlogik + Schrift erzwingen (Arial MT Pro 11)

import os  # Kommentar: OS-Pfade/Ordner
import json  # Kommentar: JSON parsing
from datetime import datetime, timedelta  # Kommentar: Datum/Frist
from docxtpl import DocxTemplate  # Kommentar: Word-Template Engine (docxtpl)
from programm_1_ki_input import KI_ANTWORT_ORDNER, BASE_DIR  # Kommentar: Pfade aus Programm 1

PROGRAMM_2_VERSION = "2025-12-26-summe-all-varianten-v2-arial11"  # Kommentar: Version zum Debuggen (siehst du in der App)

AUSGANGS_ORDNER = os.path.join(BASE_DIR, "ausgang_schreiben")  # Kommentar: Ausgabeordner
JSON_START_MARKER = "JSON_START"  # Kommentar: JSON Start Marker
JSON_END_MARKER = "JSON_END"  # Kommentar: JSON End Marker

try:  # Kommentar: Optionaler Import (falls python-docx nicht installiert ist)
    from docx import Document  # Kommentar: DOCX nachträglich öffnen/bearbeiten (python-docx)
    from docx.shared import Pt  # Kommentar: Punktgröße (11pt)
    from docx.oxml.ns import qn  # Kommentar: Word-XML Font-Felder (ascii/hAnsi/cs)
    PYDOCX_AVAILABLE = True  # Kommentar: Flag: python-docx verfügbar
except Exception:  # Kommentar: Wenn Import fehlschlägt
    Document = None  # Kommentar: Platzhalter
    Pt = None  # Kommentar: Platzhalter
    qn = None  # Kommentar: Platzhalter
    PYDOCX_AVAILABLE = False  # Kommentar: Flag: python-docx nicht verfügbar


def json_aus_ki_antwort_parsen(ki_text: str) -> dict:  # Kommentar: JSON Block aus KI Text extrahieren
    start_idx = ki_text.find(JSON_START_MARKER)  # Kommentar: Startmarker suchen
    end_idx = ki_text.find(JSON_END_MARKER)  # Kommentar: Endmarker suchen
    if start_idx == -1 or end_idx == -1:  # Kommentar: Marker fehlen?
        raise ValueError("JSON_START oder JSON_END nicht gefunden.")  # Kommentar: Fehler werfen
    json_roh = ki_text[start_idx + len(JSON_START_MARKER):end_idx].strip()  # Kommentar: Block ausschneiden
    json_roh = json_roh.replace("```json", "").replace("```", "").strip()  # Kommentar: Codefences entfernen
    first_brace = json_roh.find("{")  # Kommentar: Erste Klammer finden
    last_brace = json_roh.rfind("}")  # Kommentar: Letzte Klammer finden
    if first_brace == -1 or last_brace == -1 or first_brace >= last_brace:  # Kommentar: Ungültig?
        raise ValueError("Kein gültiger JSON-Block in KI-Antwort.")  # Kommentar: Fehler
    json_clean = json_roh[first_brace:last_brace + 1]  # Kommentar: JSON sauber ausschneiden
    json_clean = json_clean.replace(",\n}", "\n}").replace(",}", "}")  # Kommentar: Häufigen KI-Kommafehler reparieren
    return json.loads(json_clean)  # Kommentar: JSON laden und zurückgeben


def euro_zu_float(text) -> float:  # Kommentar: Euro-String robust -> float
    if isinstance(text, (int, float)):  # Kommentar: Schon Zahl?
        return float(text)  # Kommentar: Cast
    if not text:  # Kommentar: leer?
        return 0.0  # Kommentar: Null
    t = str(text)  # Kommentar: in String
    t = t.replace("€", "").replace("EUR", "").replace("Euro", "")  # Kommentar: Währung entfernen
    t = t.replace("\u00a0", " ").strip()  # Kommentar: NBSP entfernen
    t = t.replace(" ", "")  # Kommentar: Spaces entfernen
    filtered = []  # Kommentar: erlaubte Zeichen sammeln
    for ch in t:  # Kommentar: iterieren
        if ch.isdigit() or ch in [",", ".", "+", "-"]:  # Kommentar: erlaubte Zeichen
            filtered.append(ch)  # Kommentar: hinzufügen
    t = "".join(filtered)  # Kommentar: zusammenbauen
    if not t:  # Kommentar: leer?
        return 0.0  # Kommentar: Null
    if "." in t and "," in t:  # Kommentar: Fall 1.234,56
        t = t.replace(".", "").replace(",", ".")  # Kommentar: Tausender weg, Komma -> Punkt
    elif "," in t:  # Kommentar: Fall 1234,56
        t = t.replace(",", ".")  # Kommentar: Komma -> Punkt
    try:  # Kommentar: parse versuchen
        return float(t)  # Kommentar: float zurück
    except ValueError:  # Kommentar: parse fail
        return 0.0  # Kommentar: fallback


def float_zu_euro(betrag: float) -> str:  # Kommentar: float -> deutsches Euroformat
    s = f"{betrag:,.2f}"  # Kommentar: US-Format (1,234.56)
    s = s.replace(",", "X").replace(".", ",").replace("X", ".")  # Kommentar: Punkt/Komma tauschen (1.234,56)
    return s + " €"  # Kommentar: Eurozeichen anhängen


def daten_defaults(daten: dict) -> dict:  # Kommentar: Defaults für alle bekannten Keys setzen
    keys = [  # Kommentar: mögliche Template-Keys (dein Set)
        "REPARATURKOSTEN", "MWST_BETRAG", "WERTMINDERUNG", "NUTZUNGSAUSFALL",  # Kommentar: Tabelle
        "KOSTENPAUSCHALE", "GUTACHTERKOSTEN",  # Kommentar: weitere Kosten
        "WIEDERBESCHAFFUNGSWERT", "RESTWERT", "WIEDERBESCHAFFUNGSAUFWAND",  # Kommentar: Totalschaden
        "ZUSATZKOSTEN_BEZEICHNUNG", "ZUSATZKOSTEN_BETRAG",  # Kommentar: Zusatz
        "KOSTENSUMME_X", "GESAMTSUMME",  # Kommentar: Summen
        "ABRECHNUNGSART", "STEUERSTATUS", "VORSTEUERBERECHTIGUNG",  # Kommentar: Textfelder
        "FRIST_DATUM", "HEUTDATUM",  # Kommentar: Datum
        "MANDANT_VORNAME", "MANDANT_NACHNAME", "MANDANT_NAME", "MANDANT_STRASSE", "MANDANT_PLZ_ORT",  # Kommentar: Mandant
        "UNFALL_DATUM", "UNFALL_UHRZEIT", "UNFALLORT", "UNFALL_STRASSE",  # Kommentar: Unfall
        "FAHRZEUGTYP", "KENNZEICHEN", "FAHRZEUG_KENNZEICHEN",  # Kommentar: Fahrzeug
        "POLIZEIAKTE_NUMMER", "SCHADENSNUMMER", "AKTENZEICHEN",  # Kommentar: Nummern
        "SCHADENHERGANG",  # Kommentar: Hergang
        "DEBUG_PROGRAMM2_VERSION", "DEBUG_SUMME_TEILE",  # Kommentar: Debug
    ]  # Kommentar: Ende Keys
    for k in keys:  # Kommentar: iterieren
        daten.setdefault(k, "")  # Kommentar: Default setzen
    return daten  # Kommentar: Return


def setze_vorsteuer_text(daten: dict, steuerstatus: str) -> dict:  # Kommentar: setzt {{VORSTEUERBERECHTIGUNG}}
    daten["STEUERSTATUS"] = steuerstatus  # Kommentar: speichern (für Word/Debug)
    if steuerstatus == "vorsteuerabzugsberechtigt":  # Kommentar: Vorsteuer-Fall
        daten["VORSTEUERBERECHTIGUNG"] = "vorsteuerberechtigt"  # Kommentar: Text wie gewünscht
    else:  # Kommentar: Standard-Fall
        daten["VORSTEUERBERECHTIGUNG"] = "nicht vorsteuerberechtigt"  # Kommentar: Text wie gewünscht
    return daten  # Kommentar: Return


def variante_key(auswahl: str) -> str:  # Kommentar: Mappt UI-Auswahl auf internen Key
    if auswahl == "Fiktive Abrechnung (Reparaturschaden)":  # Kommentar: Fall 1
        return "FIKTIV_REPARATUR"  # Kommentar: Key
    if auswahl == "Konkrete Abrechnung < WBW":  # Kommentar: Fall 2
        return "KONKRET_UNTER_WBW"  # Kommentar: Key
    if auswahl == "130%-Regelung":  # Kommentar: Fall 3
        return "REGEL_130"  # Kommentar: Key
    if auswahl == "Totalschaden fiktiv":  # Kommentar: Fall 4
        return "TOTAL_FIKTIV"  # Kommentar: Key
    if auswahl == "Totalschaden konkret":  # Kommentar: Fall 5
        return "TOTAL_KONKRET"  # Kommentar: Key
    if auswahl == "Totalschaden Ersatzbeschaffung":  # Kommentar: Fall 6
        return "TOTAL_ERSATZ"  # Kommentar: Key
    return "UNKNOWN"  # Kommentar: Fallback


def mwst_leeren_wenn_noetig(daten: dict, auswahl: str, steuerstatus: str) -> dict:  # Kommentar: MWST_BETRAG leer lassen wo nötig
    if auswahl in ["Fiktive Abrechnung (Reparaturschaden)", "Totalschaden fiktiv"]:  # Kommentar: fiktive Fälle -> keine MwSt
        daten["MWST_BETRAG"] = ""  # Kommentar: leeren
        return daten  # Kommentar: Return
    if steuerstatus == "vorsteuerabzugsberechtigt":  # Kommentar: Vorsteuer -> MwSt wird nicht “gefordert” im Schreiben
        daten["MWST_BETRAG"] = ""  # Kommentar: leeren
        return daten  # Kommentar: Return
    return daten  # Kommentar: sonst belassen


def summe_tabelle_berechnen(  # Kommentar: Summen je Variante (entscheidet, was addiert wird)
    variant: str,  # Kommentar: interner Variant-Key
    reparatur: float,  # Kommentar: Reparaturkosten
    mwst: float,  # Kommentar: MwSt Betrag
    wertminderung: float,  # Kommentar: Wertminderung
    nutzung: float,  # Kommentar: Nutzungsausfall
    kostenpausch: float,  # Kommentar: Kostenpauschale
    gutachter: float,  # Kommentar: Gutachterkosten
    zusatz: float,  # Kommentar: Zusatzkosten
    wba: float,  # Kommentar: Wiederbeschaffungsaufwand
) -> float:  # Kommentar: Summe float
    if variant == "FIKTIV_REPARATUR":  # Kommentar: fiktiv Reparaturschaden -> ohne MwSt
        return reparatur + wertminderung + nutzung + kostenpausch + gutachter + zusatz  # Kommentar: addieren
    if variant == "KONKRET_UNTER_WBW":  # Kommentar: konkret < WBW -> MwSt wird addiert (wenn nicht geleert)
        return reparatur + mwst + wertminderung + nutzung + kostenpausch + gutachter + zusatz  # Kommentar: addieren
    if variant == "REGEL_130":  # Kommentar: 130% -> Wertminderung ausgeschlossen
        return reparatur + mwst + nutzung + kostenpausch + gutachter + zusatz  # Kommentar: addieren
    if variant == "TOTAL_FIKTIV":  # Kommentar: Totalschaden fiktiv -> WBA statt Reparatur, ohne MwSt
        return wba + nutzung + kostenpausch + gutachter + zusatz  # Kommentar: addieren
    if variant == "TOTAL_KONKRET":  # Kommentar: Totalschaden konkret -> WBA + ggf MwSt + Nebenkosten
        return wba + mwst + nutzung + kostenpausch + gutachter + zusatz  # Kommentar: addieren
    if variant == "TOTAL_ERSATZ":  # Kommentar: Ersatzbeschaffung -> WBA + MwSt (bis Grenze) + Nebenkosten
        return wba + mwst + nutzung + kostenpausch + gutachter + zusatz  # Kommentar: addieren
    return reparatur + mwst + wertminderung + nutzung + kostenpausch + gutachter + zusatz + wba  # Kommentar: Fallback


def erzwinge_schrift(docx_pfad: str, font_name: str = "Arial MT Pro", font_size_pt: int = 11) -> None:  # Kommentar: Schrift global erzwingen
    if not PYDOCX_AVAILABLE:  # Kommentar: Wenn python-docx fehlt
        return  # Kommentar: dann skip (Vorlage muss es regeln)
    doc = Document(docx_pfad)  # Kommentar: DOCX laden
    size = Pt(font_size_pt)  # Kommentar: 11pt als Word-Einheit
    if "Normal" in doc.styles:  # Kommentar: Prüfen, ob Standardstil existiert
        doc.styles["Normal"].font.name = font_name  # Kommentar: Normal-Stil Font setzen
        doc.styles["Normal"].font.size = size  # Kommentar: Normal-Stil Größe setzen
        try:  # Kommentar: XML-Fonts setzen (robuster in Word)
            doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), font_name)  # Kommentar: ASCII Font
            doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)  # Kommentar: HAnsi Font
            doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:cs"), font_name)  # Kommentar: CS Font
        except Exception:  # Kommentar: Falls XML nicht erreichbar
            pass  # Kommentar: Ignorieren

    for paragraph in doc.paragraphs:  # Kommentar: Alle Absätze
        for run in paragraph.runs:  # Kommentar: Alle Runs im Absatz
            run.font.name = font_name  # Kommentar: Run-Font setzen
            run.font.size = size  # Kommentar: Run-Größe setzen
            try:  # Kommentar: XML-Fonts setzen (robuster)
                run._element.get_or_add_rPr()  # Kommentar: rPr sicherstellen
                rfonts = run._element.rPr.rFonts  # Kommentar: rFonts holen
                if rfonts is not None:  # Kommentar: wenn vorhanden
                    rfonts.set(qn("w:ascii"), font_name)  # Kommentar: ASCII Font
                    rfonts.set(qn("w:hAnsi"), font_name)  # Kommentar: HAnsi Font
                    rfonts.set(qn("w:cs"), font_name)  # Kommentar: CS Font
            except Exception:  # Kommentar: Falls nicht möglich
                pass  # Kommentar: Ignorieren

    for table in doc.tables:  # Kommentar: Tabellen
        for row in table.rows:  # Kommentar: Zeilen
            for cell in row.cells:  # Kommentar: Zellen
                for paragraph in cell.paragraphs:  # Kommentar: Absätze in Zellen
                    for run in paragraph.runs:  # Kommentar: Runs in Zellen
                        run.font.name = font_name  # Kommentar: Font setzen
                        run.font.size = size  # Kommentar: Größe setzen
                        try:  # Kommentar: XML-Fonts setzen
                            run._element.get_or_add_rPr()  # Kommentar: rPr sicherstellen
                            rfonts = run._element.rPr.rFonts  # Kommentar: rFonts holen
                            if rfonts is not None:  # Kommentar: wenn vorhanden
                                rfonts.set(qn("w:ascii"), font_name)  # Kommentar: ASCII Font
                                rfonts.set(qn("w:hAnsi"), font_name)  # Kommentar: HAnsi Font
                                rfonts.set(qn("w:cs"), font_name)  # Kommentar: CS Font
                        except Exception:  # Kommentar: Falls nicht möglich
                            pass  # Kommentar: Ignorieren

    doc.save(docx_pfad)  # Kommentar: Speichern (überschreibt Datei)


def prepare_data_for_template(daten: dict, auswahl: str, steuerstatus: str, zus_bez: str = "", zus_betrag: str = "") -> dict:  # Kommentar: Public Funktion – bereitet Daten inkl. Summe vor
    daten = daten_defaults(daten)  # Kommentar: Defaults setzen
    daten["ABRECHNUNGSART"] = auswahl  # Kommentar: Auswahl speichern
    daten["ZUSATZKOSTEN_BEZEICHNUNG"] = (zus_bez or "").strip()  # Kommentar: Zusatzbezeichnung
    daten["ZUSATZKOSTEN_BETRAG"] = (zus_betrag or "").strip()  # Kommentar: Zusatzbetrag
    jetzt = datetime.now()  # Kommentar: jetzt
    daten["HEUTDATUM"] = jetzt.strftime("%d.%m.%Y")  # Kommentar: heutiges Datum
    daten["FRIST_DATUM"] = (jetzt + timedelta(days=14)).strftime("%d.%m.%Y")  # Kommentar: Frist
    daten = setze_vorsteuer_text(daten, steuerstatus)  # Kommentar: Vorsteuertext setzen
    daten = mwst_leeren_wenn_noetig(daten, auswahl, steuerstatus)  # Kommentar: MwSt ggf. leeren

    variant = variante_key(auswahl)  # Kommentar: interner Variant-Key

    reparatur = euro_zu_float(daten.get("REPARATURKOSTEN", ""))  # Kommentar: float Reparatur
    mwst = euro_zu_float(daten.get("MWST_BETRAG", ""))  # Kommentar: float MwSt (evtl. schon geleert)
    wertminderung = euro_zu_float(daten.get("WERTMINDERUNG", ""))  # Kommentar: float Wertminderung
    nutzung = euro_zu_float(daten.get("NUTZUNGSAUSFALL", ""))  # Kommentar: float Nutzungsausfall
    kostenpausch = euro_zu_float(daten.get("KOSTENPAUSCHALE", ""))  # Kommentar: float Kostenpauschale
    gutachter = euro_zu_float(daten.get("GUTACHTERKOSTEN", ""))  # Kommentar: float Gutachterkosten
    zusatz = euro_zu_float(daten.get("ZUSATZKOSTEN_BETRAG", ""))  # Kommentar: float Zusatzkosten

    wbw = euro_zu_float(daten.get("WIEDERBESCHAFFUNGSWERT", ""))  # Kommentar: float WBW
    rest = euro_zu_float(daten.get("RESTWERT", ""))  # Kommentar: float Restwert
    wba = max(wbw - rest, 0.0)  # Kommentar: Wiederbeschaffungsaufwand

    if variant == "REGEL_130":  # Kommentar: 130% -> Wertminderung raus
        wertminderung = 0.0  # Kommentar: auf 0 setzen
        daten["WERTMINDERUNG"] = ""  # Kommentar: im Template leeren

    if variant in ["TOTAL_FIKTIV", "TOTAL_KONKRET", "TOTAL_ERSATZ"]:  # Kommentar: Totalschaden -> Reparaturkosten nicht als Position
        reparatur = 0.0  # Kommentar: 0
        daten["REPARATURKOSTEN"] = ""  # Kommentar: Feld leeren
        if wba > 0:  # Kommentar: WBA vorhanden
            daten["WIEDERBESCHAFFUNGSAUFWAND"] = float_zu_euro(wba)  # Kommentar: Feld setzen

    summe = summe_tabelle_berechnen(variant, reparatur, mwst, wertminderung, nutzung, kostenpausch, gutachter, zusatz, wba)  # Kommentar: Summe berechnen
    daten["KOSTENSUMME_X"] = float_zu_euro(summe) if summe > 0 else ""  # Kommentar: Summe für Word
    daten["GESAMTSUMME"] = daten["KOSTENSUMME_X"]  # Kommentar: identisch (Fallback)

    if reparatur > 0:  # Kommentar: Reparatur formatieren
        daten["REPARATURKOSTEN"] = float_zu_euro(reparatur)  # Kommentar: setzen
    if mwst > 0 and daten.get("MWST_BETRAG", "") != "":  # Kommentar: MwSt nur formatieren, wenn nicht geleert
        daten["MWST_BETRAG"] = float_zu_euro(mwst)  # Kommentar: setzen
    if wertminderung > 0:  # Kommentar: Wertminderung formatieren
        daten["WERTMINDERUNG"] = float_zu_euro(wertminderung)  # Kommentar: setzen
    if nutzung > 0:  # Kommentar: Nutzung formatieren
        daten["NUTZUNGSAUSFALL"] = float_zu_euro(nutzung)  # Kommentar: setzen
    if kostenpausch > 0:  # Kommentar: Pauschale formatieren
        daten["KOSTENPAUSCHALE"] = float_zu_euro(kostenpausch)  # Kommentar: setzen
    if gutachter > 0:  # Kommentar: Gutachter formatieren
        daten["GUTACHTERKOSTEN"] = float_zu_euro(gutachter)  # Kommentar: setzen
    if zusatz > 0:  # Kommentar: Zusatz formatieren
        daten["ZUSATZKOSTEN_BETRAG"] = float_zu_euro(zusatz)  # Kommentar: setzen

    daten["DEBUG_PROGRAMM2_VERSION"] = PROGRAMM_2_VERSION  # Kommentar: Debug: Version
    daten["DEBUG_SUMME_TEILE"] = (  # Kommentar: Debugstring
        f"variant={variant} reparatur={reparatur} mwst={mwst} wm={wertminderung} "
        f"nutzung={nutzung} pausch={kostenpausch} gutachter={gutachter} zusatz={zusatz} "
        f"wba={wba} => summe={summe}"
    )  # Kommentar: Ende Debugstring

    return daten  # Kommentar: Return


def word_aus_vorlage_erstellen(daten: dict, vorlage_pfad: str, ziel_pfad: str) -> None:  # Kommentar: DOCX rendern
    doc = DocxTemplate(vorlage_pfad)  # Kommentar: Template laden
    os.makedirs(os.path.dirname(ziel_pfad), exist_ok=True)  # Kommentar: Ordner sicherstellen
    doc.render(daten)  # Kommentar: rendern (alle Keys verfügbar)
    doc.save(ziel_pfad)  # Kommentar: speichern
    erzwinge_schrift(ziel_pfad, "Arial MT Pro", 11)  # Kommentar: Schrift global erzwingen (wenn python-docx verfügbar)


def ki_datei_verarbeiten(pfad_ki_txt: str, vorlage_pfad: str, auswahl: str, steuerstatus: str, zus_bez: str, zus_betrag: str) -> str:  # Kommentar: KI-Datei -> DOCX
    with open(pfad_ki_txt, "r", encoding="utf-8") as f:  # Kommentar: Datei öffnen
        ki_text = f.read()  # Kommentar: lesen
    daten = json_aus_ki_antwort_parsen(ki_text)  # Kommentar: JSON parsen
    daten = prepare_data_for_template(daten, auswahl, steuerstatus, zus_bez, zus_betrag)  # Kommentar: nachbearbeiten + Summe
    datum_str = datetime.now().strftime("%Y%m%d_%H%M%S")  # Kommentar: Timestamp
    out_name = f"schreiben_{datum_str}.docx"  # Kommentar: Ausgabename
    out_path = os.path.join(AUSGANGS_ORDNER, out_name)  # Kommentar: Ausgabe-Pfad
    word_aus_vorlage_erstellen(daten, vorlage_pfad, out_path)  # Kommentar: rendern/speichern
    return out_path  # Kommentar: Pfad zurückgeben


def main(pfad_ki_txt: str = None, vorlage_pfad: str | None = None, auswahl: str = "", steuerstatus: str = "", zus_bez: str = "", zus_betrag: str = "") -> str:  # Kommentar: Main für App
    os.makedirs(KI_ANTWORT_ORDNER, exist_ok=True)  # Kommentar: Ordner sicher
    os.makedirs(AUSGANGS_ORDNER, exist_ok=True)  # Kommentar: Ordner sicher
    if vorlage_pfad is None:  # Kommentar: Vorlage nötig
        raise ValueError("vorlage_pfad muss übergeben werden (Word-Vorlage .docx).")  # Kommentar: Fehler
    if pfad_ki_txt is None:  # Kommentar: wenn nicht gegeben
        files = [os.path.join(KI_ANTWORT_ORDNER, f) for f in os.listdir(KI_ANTWORT_ORDNER) if f.endswith("_ki.txt")]  # Kommentar: suchen
        if not files:  # Kommentar: keine gefunden?
            raise FileNotFoundError("Keine KI-Datei gefunden.")  # Kommentar: Fehler
        pfad_ki_txt = max(files, key=os.path.getmtime)  # Kommentar: neueste nehmen
    if not os.path.isfile(pfad_ki_txt):  # Kommentar: existiert?
        raise FileNotFoundError(f"KI-Datei existiert nicht: {pfad_ki_txt}")  # Kommentar: Fehler
    if not os.path.isfile(vorlage_pfad):  # Kommentar: Vorlage existiert?
        raise FileNotFoundError(f"Vorlage existiert nicht: {vorlage_pfad}")  # Kommentar: Fehler
    return ki_datei_verarbeiten(pfad_ki_txt, vorlage_pfad, auswahl, steuerstatus, zus_bez, zus_betrag)  # Kommentar: run


if __name__ == "__main__":  # Kommentar: Direktstart
    main()  # Kommentar: aufrufen
